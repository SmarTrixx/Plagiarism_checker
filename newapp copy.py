import os
import pytesseract
from pdf2image import convert_from_path
from PIL import Image
from docx import Document
from sklearn.metrics.pairwise import cosine_similarity
from sentence_transformers import SentenceTransformer
from tqdm import tqdm
import tempfile

# ========== Config ==========
FOLDER_PATH = "/home/smartz/Desktop/Major Projects/Plagiarism Checker/uploads/archive"  # Folder with all files
SIMILARITY_THRESHOLD = 0.5   # Customize as needed
TOP_RESULTS = 10             # Show top N similar pairs
OCR_LANG = 'eng'             # Change if using another language
MODEL_NAME = 'all-MiniLM-L6-v2'
# =============================

# Load Sentence-BERT model
model = SentenceTransformer(MODEL_NAME)

# Helper: extract text from DOCX
def extract_docx(file_path):
    doc = Document(file_path)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

# Helper: extract text from image (OCR)
def extract_image_text(img):
    img = img.convert('L')  # grayscale
    return pytesseract.image_to_string(img, lang=OCR_LANG)

# Helper: extract text from PDF (OCR page-by-page)
def extract_pdf(file_path):
    text = ""
    images = convert_from_path(file_path)
    for img in images:
        text += extract_image_text(img) + "\n"
    return text

# Extract plain text from file
def extract_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext == ".docx":
            return extract_docx(file_path)
        elif ext == ".pdf":
            return extract_pdf(file_path)
        elif ext in [".png", ".jpg", ".jpeg"]:
            return extract_image_text(Image.open(file_path))
        else:
            return ""  # Unsupported
    except Exception as e:
        print(f"[ERROR] {file_path}: {e}")
        return ""

# Get all supported files from folder
def get_all_files(folder):
    supported_ext = [".docx", ".pdf", ".jpg", ".jpeg", ".png"]
    return [os.path.join(folder, f) for f in os.listdir(folder)
            if os.path.splitext(f)[1].lower() in supported_ext]

# Compare text using sentence-BERT + cosine similarity
def compare_texts(text1, text2):
    sentences1 = [s.strip() for s in text1.split(".") if len(s.strip()) > 10]
    sentences2 = [s.strip() for s in text2.split(".") if len(s.strip()) > 10]
    if not sentences1 or not sentences2:
        return 0.0

    emb1 = model.encode(sentences1, convert_to_tensor=True)
    emb2 = model.encode(sentences2, convert_to_tensor=True)

    sim_matrix = cosine_similarity(emb1, emb2)
    max_sim_per_row = sim_matrix.max(axis=1)
    avg_sim = max_sim_per_row.mean()

    return round(float(avg_sim), 4)

# Main comparison loop
def main():
    files = get_all_files(FOLDER_PATH)
    results = []

    print(f"📂 Found {len(files)} files. Extracting texts...")

    file_texts = {}
    for f in tqdm(files, desc="Extracting"):
        file_texts[f] = extract_text(f)

    print("🔍 Comparing files... (this may take a while)")
    for i in tqdm(range(len(files))):
        for j in range(i+1, len(files)):
            f1, f2 = files[i], files[j]
            sim = compare_texts(file_texts[f1], file_texts[f2])
            if sim >= SIMILARITY_THRESHOLD:
                results.append((os.path.basename(f1), os.path.basename(f2), sim))

    if not results:
        print("✅ No significant similarities found.")
    else:
        print(f"\n🔗 Top {TOP_RESULTS} Similar Documents (Threshold: {SIMILARITY_THRESHOLD}):\n")
        results.sort(key=lambda x: x[2], reverse=True)
        for r in results[:TOP_RESULTS]:
            print(f"{r[0]}  🔁  {r[1]}  |  Similarity: {r[2]*100:.2f}%")

if __name__ == "__main__":
    main()






//////////////////////

import os
import pytesseract
import docx
import fitz  # PyMuPDF
from PIL import Image
from pdf2image import convert_from_path
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np

# === Configurations ===
TESSERACT_PATH = r"/usr/bin/tesseract"  # Set this if not auto-detected
pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH
model = SentenceTransformer("all-MiniLM-L6-v2")
SIMILARITY_THRESHOLD = 0.85  # adjust as needed

# === Preprocessing for OCR ===
def extract_image_text(image):
    gray = image.convert('L')
    return pytesseract.image_to_string(gray)

# === DOCX Extract ===
def extract_docx(file_path):
    doc = docx.Document(file_path)
    return '\n'.join([para.text for para in doc.paragraphs])

# === Image File Extract ===
def extract_image(file_path):
    img = Image.open(file_path)
    return extract_image_text(img)

# === Smart PDF Extract (Typed or OCR fallback) ===
def extract_pdf(file_path):
    try:
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        if text.strip():
            return text  # ✅ Typed PDF handled
    except Exception as e:
        print(f"[Typed PDF failed] {file_path}: {e}")

    # Fallback to OCR
    try:
        text = ""
        images = convert_from_path(file_path)
        for img in images:
            text += extract_image_text(img) + "\n"
        return text
    except Exception as e:
        print(f"[OCR PDF failed] {file_path}: {e}")
        return ""

# === File Dispatcher ===
def extract_text(file_path):
    ext = file_path.lower()
    if ext.endswith('.pdf'):
        return extract_pdf(file_path)
    elif ext.endswith('.docx'):
        return extract_docx(file_path)
    elif ext.endswith(('.png', '.jpg', '.jpeg')):
        return extract_image(file_path)
    else:
        print(f"Unsupported file format: {file_path}")
        return ""

# === Sentence Comparison ===
def compare_sentences(text1, text2):
    sents1 = [s.strip() for s in text1.split('.') if len(s.strip()) > 10]
    sents2 = [s.strip() for s in text2.split('.') if len(s.strip()) > 10]

    if not sents1 or not sents2:
        return []  # 🔒 Avoid comparing if one side has no valid sentences

    try:
        emb1 = model.encode(sents1, convert_to_tensor=True)
        emb2 = model.encode(sents2, convert_to_tensor=True)

        sim_matrix = cosine_similarity(emb1, emb2)
        similar_pairs = []

        for i, row in enumerate(sim_matrix):
            for j, sim in enumerate(row):
                if sim > SIMILARITY_THRESHOLD:
                    similar_pairs.append((sents1[i], sents2[j], sim))

        return similar_pairs

    except Exception as e:
        print(f"[Error comparing sentences] {e}")
        if not sents1:
            print(f"[Skipped] No valid sentences in: {file1}")
        if not sents2:
            print(f"[Skipped] No valid sentences in: {file2}")

        return []

    sents1 = [s.strip() for s in text1.split('.') if len(s.strip()) > 10]
    sents2 = [s.strip() for s in text2.split('.') if len(s.strip()) > 10]
    emb1 = model.encode(sents1, convert_to_tensor=True)
    emb2 = model.encode(sents2, convert_to_tensor=True)

    sim_matrix = cosine_similarity(emb1, emb2)
    similar_pairs = []

    for i, row in enumerate(sim_matrix):
        for j, sim in enumerate(row):
            if sim > SIMILARITY_THRESHOLD:
                similar_pairs.append((sents1[i], sents2[j], sim))

    return similar_pairs

# === Main Runner ===
def process_folder(folder_path):
    files = [os.path.join(folder_path, f) for f in os.listdir(folder_path)
             if f.lower().endswith(('.pdf', '.docx', '.jpg', '.jpeg', '.png'))]
    texts = {}
    for file in files:
        print(f"Extracting: {file}")
        texts[file] = extract_text(file)

    checked = set()
    print("\n=== Possible Plagiarism Detected ===\n")
    for i, file1 in enumerate(files):
        for j, file2 in enumerate(files):
            if i >= j or (file1, file2) in checked or (file2, file1) in checked:
                continue
            sim_pairs = compare_sentences(texts[file1], texts[file2])
            if sim_pairs:
                print(f"\n📁 {os.path.basename(file1)} vs {os.path.basename(file2)}")
                for s1, s2, score in sim_pairs[:3]:  # show top 3 similar
                    print(f"  - Sim: {score:.2f}")
                    print(f"    ↪ \"{s1.strip()}\"")
                    print(f"    ↪ \"{s2.strip()}\"\n")
            checked.add((file1, file2))

# === Run ===
if __name__ == "__main__":
    folder_to_check = "/home/smartz/Desktop/Major Projects/Plagiarism Checker/uploads/archive"  # <== CHANGE THIS
    process_folder(folder_to_check)
